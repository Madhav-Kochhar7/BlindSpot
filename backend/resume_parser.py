"""
resume_parser.py

Utilities for extracting plain text from resume files (.pdf and .docx).
"""

from __future__ import annotations

import io
import os

import pdfplumber
from docx import Document


class UnsupportedFileTypeError(Exception):
    """Raised when a file extension is not supported by the parser."""


class ResumeParsingError(Exception):
    """Raised when a file cannot be parsed (corrupted, empty, image-only, etc.)."""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF file's raw bytes using pdfplumber."""
    text_chunks: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text:
                    text_chunks.append(page_text)
    except Exception as exc:  # pdfplumber can raise several exception types
        raise ResumeParsingError(f"Failed to parse PDF: {exc}") from exc

    return "\n".join(text_chunks).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX file's raw bytes using python-docx."""
    text_chunks: list[str] = []
    try:
        document = Document(io.BytesIO(file_bytes))

        for paragraph in document.paragraphs:
            if paragraph.text:
                text_chunks.append(paragraph.text)

        # Tables often hold skills/experience blocks in resumes - capture them too.
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_chunks.append(cell.text)
    except Exception as exc:
        raise ResumeParsingError(f"Failed to parse DOCX: {exc}") from exc

    return "\n".join(text_chunks).strip()


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatch to the correct extractor based on file extension.

    Raises:
        UnsupportedFileTypeError: if the extension isn't .pdf or .docx
        ResumeParsingError: if the file can't be parsed or has no text
    """
    extension = os.path.splitext(filename)[1].lower()

    if extension == ".pdf":
        text = extract_text_from_pdf(file_bytes)
    elif extension == ".docx":
        text = extract_text_from_docx(file_bytes)
    else:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{extension}'. Only .pdf and .docx are supported."
        )

    if not text:
        raise ResumeParsingError(
            "No extractable text found. The file may be a scanned image or empty."
        )

    return text
