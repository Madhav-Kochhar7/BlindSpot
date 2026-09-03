"""
test_phase2_e2e.py

End-to-end verification script for Phase 2 dual pipeline.
Tests redaction, evidence extraction, rank calculation, and API output.
"""

import asyncio
import io
from redaction import redact_resume
from evidence_extractor import extract_and_score_evidence
from ats_scorer import extract_keywords, score_resume
from main import app
from fastapi.testclient import TestClient

SAMPLE_JD = """
Senior Full-Stack Engineer

Requirements:
- 4+ years of experience with React, TypeScript, and modern state management.
- Backend proficiency with Python, FastAPI, and PostgreSQL.
- Experience with Docker, microservices, and CI/CD pipelines.
- Track record of quantifiable outcomes (e.g. improved performance by 40%, scaled to 100k users).
"""

# Candidate 1: Genuine high-evidence candidate from a lesser-known college
RESUME_1 = """
Alex Rivera
Email: alex.rivera@gmail.com
Phone: (415) 555-0199
B.S. Computer Science from Portland State University, 2017-2021.
Portfolio: https://alexrivera.dev

Professional Summary:
Lead Full-Stack Developer with 4 years of experience building resilient cloud applications.

Experience:
Senior Software Engineer (2021 - Present)
- Architected and built a high-throughput dashboard using React, TypeScript, and FastAPI that served 80,000 daily active users with 99.95% uptime.
- Optimized PostgreSQL database queries and indexing, reducing API response latency by 45% (from 420ms to 230ms).
- Containerized microservices using Docker and established automated CI/CD pipelines deploying to AWS ECS.
- Designed secure RESTful APIs with Python and Pydantic validation handling over 5,000 req/s.
"""

# Candidate 2: Keyword-stuffed resume from a prestigious college with weak proof
RESUME_2 = """
Bradford Sterling III
Email: b.sterling@stanford.edu
Phone: +1 650-555-8888
B.S. and M.S. in Computer Science from Stanford University (2015-2019).
LinkedIn: linkedin.com/in/bsterling

Skills:
React, TypeScript, Python, FastAPI, PostgreSQL, Docker, CI/CD, AWS, Microservices, SQL, Kubernetes, Redis, GraphQL, Agile, Leadership, Innovation, Scalability.

Experience:
Software Consultant (2019 - Present)
- Worked with various clients on React and TypeScript web apps.
- Familiar with Python, FastAPI, and backend development.
- Interacted with PostgreSQL databases and Docker containers.
- Attended weekly sprint planning and team meetings.
"""

def test_redaction():
    print("\n--- Testing PII Redaction ---")
    anon_id_1, redacted_1, stats_1 = redact_resume(RESUME_1, candidate_index=1, candidate_name="Alex Rivera")
    print(f"Candidate 1 Anon ID: {anon_id_1}")
    print(f"Candidate 1 Stats: {stats_1}")
    assert "Alex Rivera" not in redacted_1
    assert "alex.rivera@gmail.com" not in redacted_1
    assert "[REDACTED_INSTITUTION]" in redacted_1
    assert "[REDACTED_DATE]" in redacted_1

    anon_id_2, redacted_2, stats_2 = redact_resume(RESUME_2, candidate_index=2, candidate_name="Bradford Sterling III")
    print(f"Candidate 2 Anon ID: {anon_id_2}")
    print(f"Candidate 2 Stats: {stats_2}")
    assert "Stanford" not in redacted_2
    assert "Bradford" not in redacted_2
    print("[PASS] PII Redaction tests passed successfully!")

def test_evidence_scoring():
    print("\n--- Testing Evidence Scoring ---")
    anon_id_1, redacted_1, _ = redact_resume(RESUME_1, candidate_index=1, candidate_name="Alex Rivera")
    anon_id_2, redacted_2, _ = redact_resume(RESUME_2, candidate_index=2, candidate_name="Bradford Sterling III")

    res_1 = asyncio.run(extract_and_score_evidence(redacted_1, SAMPLE_JD, anon_id_1))
    res_2 = asyncio.run(extract_and_score_evidence(redacted_2, SAMPLE_JD, anon_id_2))

    print(f"Candidate 1 (High Evidence) Evidence Score: {res_1.evidence_score}")
    print(f"Candidate 1 Criteria: {res_1.criteria_breakdown}")
    print(f"Candidate 2 (Keyword Stuffed) Evidence Score: {res_2.evidence_score}")
    print(f"Candidate 2 Criteria: {res_2.criteria_breakdown}")

    # Verify that Candidate 1 scores higher in evidence than Candidate 2
    assert res_1.evidence_score > res_2.evidence_score
    print("[PASS] Evidence Scoring correctly rewarded verifiable proof over keyword listing!")

def test_api_upload_and_score():
    print("\n--- Testing /api/upload-and-score with Docx Resumes ---")
    from docx import Document
    client = TestClient(app)

    # Generate real docx file in memory for Alex Rivera
    doc1 = Document()
    for line in RESUME_1.strip().split("\n"):
        doc1.add_paragraph(line)
    bio1 = io.BytesIO()
    doc1.save(bio1)
    bio1.seek(0)

    # Generate real docx file in memory for Bradford Sterling
    doc2 = Document()
    for line in RESUME_2.strip().split("\n"):
        doc2.add_paragraph(line)
    bio2 = io.BytesIO()
    doc2.save(bio2)
    bio2.seek(0)

    response = client.post(
        "/api/upload-and-score",
        data={"job_description": SAMPLE_JD},
        files=[
            ("resumes", ("alex_rivera_resume.docx", bio1.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
            ("resumes", ("bradford_sterling_resume.docx", bio2.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")),
        ],
    )

    assert response.status_code == 200, response.text
    data = response.json()
    print("Total resumes scored:", data["total_resumes"])
    print("Candidates in response:", len(data["candidates"]))
    for c in data["candidates"]:
        print(f"Candidate: {c['anonymized_id']} ({c['candidate_name']}) | ATS Score: {c['ats_score']}% (Rank #{c['rank_ats']}) | Evidence Score: {c['evidence_score']}% (Rank #{c['rank_evidence']}) | Delta: {c['rank_delta']}")

def test_sample_resumes_batch():
    print("\n--- Testing Batch with Generated Sample Resumes ---")
    import glob
    import os
    client = TestClient(app)

    files_list = []
    for filepath in sorted(glob.glob("../sample_resumes/*.docx")):
        filename = os.path.basename(filepath)
        with open(filepath, "rb") as f:
            files_list.append(
                ("resumes", (filename, f.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
            )

    response = client.post(
        "/api/upload-and-score",
        data={"job_description": SAMPLE_JD},
        files=files_list,
    )
    assert response.status_code == 200
    data = response.json()
    print(f"Evaluated {data['total_resumes']} candidates in batch:")
    for c in data["candidates"]:
        print(f"  {c['anonymized_id']} ({c['candidate_name']}): ATS {c['ats_score']}% (Rank #{c['rank_ats']}) | Evidence {c['evidence_score']}% (Rank #{c['rank_evidence']}) | Delta: {c['rank_delta']:+d}")
    print("[PASS] Batch sample resumes evaluated with full provenance and ranking deltas!")

if __name__ == "__main__":
    test_redaction()
    test_evidence_scoring()
    test_api_upload_and_score()
    test_sample_resumes_batch()
    print("\n==========================================")
    print("ALL PHASE 2 BACKEND TESTS PASSED!")
    print("==========================================")
